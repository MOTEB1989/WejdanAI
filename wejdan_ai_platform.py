#!/usr/bin/env python3
import argparse
import datetime
import hashlib
import json
import os
import re
import time
from typing import List, Optional, Tuple

import pandas as pd
import requests
from docx import Document

# ------------------------------------------------------------------------------------
# ------------------------ Configuration (Env / Constants) ---------------------------
# ------------------------------------------------------------------------------------

# Notion config
NOTION_TOKEN = os.getenv("NOTION_TOKEN", "")
DATABASE_ID = os.getenv("DATABASE_ID", "")
NOTION_VERSION = os.getenv("NOTION_VERSION", "2022-06-28")

NOTION_HEADERS = {
    "Authorization": f"Bearer {NOTION_TOKEN}" if NOTION_TOKEN else "",
    "Content-Type": "application/json",
    "Notion-Version": NOTION_VERSION,
}

# AI Models & KB
SUPPORTED_MODELS = [
    "gpt",
    "gemini",
    "claude",
    "perplexity",
    "kimi",
    "taskade",
    "qwen",
    "copilot",
]

# ------------------------------------------------------------------------------------
# ------------------------------ Helpers & Utils ------------------------------------
# ------------------------------------------------------------------------------------


def _sha256(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def request_with_retry(
    method: str,
    url: str,
    *,
    headers: Optional[dict] = None,
    json_body: Optional[dict] = None,
    max_retries: int = 5,
) -> Optional[requests.Response]:
    headers = headers or {}
    last_exception: Optional[Exception] = None
    for attempt in range(max_retries):
        try:
            resp = requests.request(method, url, headers=headers, json=json_body, timeout=30)
            if resp.status_code in (200, 201):
                return resp
            if resp.status_code in (429,) or 500 <= resp.status_code <= 599:
                time.sleep(1 * (2**attempt))
                continue
            return resp
        except Exception as exc:  # noqa: BLE001 - surface retry behavior
            last_exception = exc
            time.sleep(1)
    if last_exception:
        print(f"❌ Request failed: {last_exception}")
    return None


# ------------------------------------------------------------------------------------
# --------------------------- BSM Excel & Word Reporting -----------------------------
# ------------------------------------------------------------------------------------


class BSMProcessor:
    def __init__(self, path: str):
        self.path = path
        self.df: Optional[pd.DataFrame] = None

    def load_data(self) -> None:
        try:
            self.df = pd.read_excel(self.path, engine="openpyxl")
            print(f"Loaded BSM Excel: {len(self.df)} rows")
        except Exception as exc:  # noqa: BLE001 - surface failure context
            raise RuntimeError(f"Failed to load Excel: {exc}") from exc

    def calculate_fte(self) -> float:
        if self.df is None:
            raise RuntimeError("Data not loaded")

        if "Annual Hours" in self.df.columns:
            self.df["FTE_Required"] = self.df["Annual Hours"] / 2000
            return float(self.df["FTE_Required"].sum())
        return 0.0

    def generate_word(self, output: str = "BSM_Report.docx") -> str:
        if self.df is None:
            self.load_data()

        doc = Document()
        doc.add_heading("BSM Strategic Report", 0)
        total_fte = self.calculate_fte()
        doc.add_paragraph(f"Total FTE Required: {total_fte:.2f}")
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Theme"
        hdr[1].text = "Task"
        hdr[2].text = "Hours"
        if self.df is not None and "Annual Hours" in self.df.columns:
            for _, row in self.df.head(10).iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row.get("Theme", ""))
                cells[1].text = str(row.get("Task_Description", ""))
                cells[2].text = str(row.get("Annual Hours", ""))
        doc.save(output)
        return output


# ------------------------------------------------------------------------------------
# --------------------------- ChatGPT Shared / Ingestion Logic -----------------------
# ------------------------------------------------------------------------------------


def fetch_shared_chat(url: str) -> Tuple[str, str]:
    try:
        resp = requests.get(url, timeout=15)
        if resp.status_code != 200:
            return "", f"Fetch failed {resp.status_code}"
        text = re.sub(r"<[^>]+>", "\n", resp.text)
        return text.strip(), ""
    except Exception as exc:  # noqa: BLE001 - return error string
        return "", str(exc)


def load_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8") as file:
        return file.read()


# ------------------------------------------------------------------------------------
# ----------------------------- Knowledge Base (Flat Files) -------------------------
# ------------------------------------------------------------------------------------


class AIDatabase:
    def __init__(self, root: str = "ai_knowledge_base"):
        self.root = root
        os.makedirs(self.root, exist_ok=True)
        for model in SUPPORTED_MODELS:
            os.makedirs(os.path.join(self.root, model), exist_ok=True)
        self.index_file = os.path.join(self.root, "db_index.json")

    def save_md(
        self,
        model: str,
        title: str,
        content: str,
        url: str = "",
        tags: Optional[List[str]] = None,
    ) -> str:
        model = model.lower()
        if model not in SUPPORTED_MODELS:
            raise ValueError(f"Unsupported model: {SUPPORTED_MODELS}")

        safe_title = re.sub(r"[\\/*?\"<>|]", "", title).replace(" ", "_")[:50]
        stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{stamp}_{safe_title}.md"
        path = os.path.join(self.root, model, filename)

        metadata = {
            "id": stamp,
            "title": title,
            "model": model,
            "url": url,
            "tags": tags or [],
            "date": datetime.datetime.now().isoformat(),
        }
        with open(path, "w", encoding="utf-8") as file:
            file.write("---\n")
            file.write(json.dumps(metadata, ensure_ascii=False, indent=2))
            file.write("\n---\n\n")
            file.write(content)

        self._update_index()
        return path

    def _update_index(self) -> None:
        entries = []
        for model in SUPPORTED_MODELS:
            md_dir = os.path.join(self.root, model)
            for fname in os.listdir(md_dir):
                if fname.endswith(".md"):
                    full = os.path.join(md_dir, fname)
                    with open(full, "r", encoding="utf-8") as file:
                        body = file.read()
                        if body.startswith("---"):
                            end = body.find("---", 3)
                            if end != -1:
                                meta = json.loads(body[3:end].strip())
                                meta["filename"] = f"{model}/{fname}"
                                entries.append(meta)
        with open(self.index_file, "w", encoding="utf-8") as file:
            json.dump(entries, file, ensure_ascii=False, indent=2)


# ------------------------------------------------------------------------------------
# ----------------------------- Notion Integration ----------------------------------
# ------------------------------------------------------------------------------------


def notion_upsert(title: str, content_snippet: str, external_id: str, model: str = "GPT"):
    if not NOTION_TOKEN or not DATABASE_ID:
        return False, "Missing Notion config"

    exists, page_id = notion_find(external_id)
    payload = {
        "parent": {"database_id": DATABASE_ID},
        "properties": {
            "Page Title": {"title": [{"text": {"content": title}}]},
            "Status": {"status": {"name": "Done"}},
            "External ID": {"rich_text": [{"text": {"content": external_id}}]},
            "Category": {"select": {"name": model}},
            "Conversation Content": {"rich_text": [{"text": {"content": content_snippet}}]},
        },
    }

    if exists and page_id:
        url = f"https://api.notion.com/v1/pages/{page_id}"
        resp = request_with_retry("PATCH", url, headers=NOTION_HEADERS, json_body=payload)
        return (resp is not None and resp.status_code == 200), getattr(resp, "text", "")
    url = "https://api.notion.com/v1/pages"
    resp = request_with_retry("POST", url, headers=NOTION_HEADERS, json_body=payload)
    return (resp is not None and resp.status_code in (200, 201)), getattr(resp, "text", "")


def notion_find(external_id: str) -> Tuple[bool, Optional[str]]:
    query_url = f"https://api.notion.com/v1/databases/{DATABASE_ID}/query"
    body = {"filter": {"property": "External ID", "rich_text": {"equals": external_id}}}
    resp = request_with_retry("POST", query_url, headers=NOTION_HEADERS, json_body=body)
    if resp is not None and resp.status_code == 200:
        results = resp.json().get("results", [])
        if results:
            return True, results[0]["id"]
    return False, None


# ------------------------------------------------------------------------------------
# ---------------------------------- CLI --------------------------------------------
# ------------------------------------------------------------------------------------


def cmd_chat(args: argparse.Namespace) -> None:
    content = ""
    note = ""
    if args.url:
        content, note = fetch_shared_chat(args.url)
    if args.file and not content:
        content = load_text_file(args.file)
    if not content:
        print("No content found:", note)
        return

    db = AIDatabase()
    tags = args.tags.split(",") if args.tags else []
    path = db.save_md(args.model, args.title, content, args.url or "", tags)
    print("Archived to:", path)

    if args.push_notion:
        snippet = content[:1800]
        eid = _sha256(args.url or args.title + snippet)
        ok, msg = notion_upsert(args.title, snippet, eid, args.model)
        print("Notion:", ok, msg)


def cmd_search(args: argparse.Namespace) -> None:
    idx = "ai_knowledge_base/db_index.json"
    if not os.path.exists(idx):
        print("No index found")
        return
    with open(idx, "r", encoding="utf-8") as file:
        data = json.load(file)
    needle = args.query.lower()
    res = [
        entry
        for entry in data
        if needle in (
            entry.get("title", "").lower() + " " + " ".join(entry.get("tags", [])).lower()
        )
    ]
    for entry in res:
        print(
            f"{entry['model']} | {entry['title']} | tags:{entry.get('tags')} | {entry.get('filename')}"
        )


def main() -> None:
    parser = argparse.ArgumentParser(prog="wejdan_cli")
    sub = parser.add_subparsers(dest="command", required=True)

    # BSM
    p_bsm = sub.add_parser("bsm", help="Analyze BSM Excel")
    p_bsm.add_argument("--excel", required=True, help="Path to BSM Excel")
    p_bsm.add_argument("--out", default="BSM_Report.docx")
    p_bsm.set_defaults(
        func=lambda args: print("Saved:", BSMProcessor(args.excel).generate_word(args.out))
    )

    # Chat ingest
    p_chat = sub.add_parser("chat", help="Ingest Chat (URL or file)")
    p_chat.add_argument("--url", help="Shared chat URL")
    p_chat.add_argument("--file", help="Fallback local text file")
    p_chat.add_argument("--model", default="gpt", choices=SUPPORTED_MODELS)
    p_chat.add_argument("--title", required=True)
    p_chat.add_argument("--tags", help="Comma tags")
    p_chat.add_argument("--push-notion", action="store_true")
    p_chat.set_defaults(func=cmd_chat)

    # Search KB
    p_search = sub.add_parser("search", help="Search Knowledge Base")
    p_search.add_argument("--query", help="Text to search in titles/tags")
    p_search.set_defaults(func=cmd_search)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
